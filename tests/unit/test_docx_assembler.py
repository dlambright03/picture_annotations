"""
Unit tests for DOCX document assembler.

Tests alt-text application to DOCX documents with position
preservation and error handling.
"""

from datetime import datetime
from pathlib import Path
from unittest.mock import MagicMock, Mock, patch

import pytest
from docx import Document
from docx.oxml.ns import qn

from ada_annotator.document_processors import DOCXAssembler
from ada_annotator.exceptions import ProcessingError
from ada_annotator.models import AltTextResult


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    """Create a sample DOCX file for testing."""
    docx_path = tmp_path / "test.docx"

    # Create document with a paragraph
    doc = Document()
    doc.add_paragraph("Test paragraph with image")
    doc.save(str(docx_path))

    return docx_path


@pytest.fixture
def output_path(tmp_path: Path) -> Path:
    """Create output path for testing."""
    return tmp_path / "output" / "test_output.docx"


@pytest.fixture
def sample_alt_text_result() -> AltTextResult:
    """Create sample alt-text result."""
    return AltTextResult(
        image_id="img-0-0",  # paragraph 0, image 0
        alt_text="A diagram showing data flow between components.",
        confidence_score=0.92,
        validation_passed=True,
        validation_warnings=[],
        tokens_used=150,
        processing_time_seconds=1.5,
        timestamp=datetime.now(),
    )


class TestDOCXAssemblerInitialization:
    """Test DOCXAssembler initialization."""

    def test_initialization_success(
        self,
        sample_docx: Path,
        output_path: Path
    ):
        """Test successful initialization."""
        assembler = DOCXAssembler(sample_docx, output_path)

        assert assembler.input_path == sample_docx
        assert assembler.output_path == output_path
        assert assembler.document is not None
        assert assembler.get_document_format() == "DOCX"

    def test_initialization_creates_output_directory(
        self,
        sample_docx: Path,
        output_path: Path
    ):
        """Test output directory is created."""
        assert not output_path.parent.exists()

        DOCXAssembler(sample_docx, output_path)

        assert output_path.parent.exists()

    def test_initialization_file_not_found(
        self,
        tmp_path: Path,
        output_path: Path
    ):
        """Test initialization with non-existent file."""
        non_existent = tmp_path / "does_not_exist.docx"

        with pytest.raises(FileNotFoundError) as exc_info:
            DOCXAssembler(non_existent, output_path)

        assert "not found" in str(exc_info.value).lower()

    def test_initialization_invalid_format(
        self,
        tmp_path: Path,
        output_path: Path
    ):
        """Test initialization with non-DOCX file."""
        txt_file = tmp_path / "test.txt"
        txt_file.write_text("Not a DOCX file")

        with pytest.raises(ValueError) as exc_info:
            DOCXAssembler(txt_file, output_path)

        assert "not a docx file" in str(exc_info.value).lower()

    def test_initialization_corrupted_docx(
        self,
        tmp_path: Path,
        output_path: Path
    ):
        """Test initialization with corrupted DOCX."""
        corrupted = tmp_path / "corrupted.docx"
        corrupted.write_bytes(b"This is not a valid DOCX file")

        with pytest.raises(ProcessingError) as exc_info:
            DOCXAssembler(corrupted, output_path)

        assert "failed to load" in str(exc_info.value).lower()


class TestDOCXAssemblerAltTextApplication:
    """Test alt-text application functionality."""

    def test_apply_alt_text_success(
        self,
        sample_docx: Path,
        output_path: Path,
        sample_alt_text_result: AltTextResult
    ):
        """Test successful alt-text application."""
        assembler = DOCXAssembler(sample_docx, output_path)

        # Mock the internal method to simulate success
        assembler._apply_alt_text_to_image = Mock(
            return_value="success"
        )

        status_map = assembler.apply_alt_text([sample_alt_text_result])

        assert len(status_map) == 1
        assert status_map[sample_alt_text_result.image_id] == "success"

    def test_apply_alt_text_multiple_results(
        self,
        sample_docx: Path,
        output_path: Path
    ):
        """Test alt-text application with multiple results."""
        assembler = DOCXAssembler(sample_docx, output_path)

        results = [
            AltTextResult(
                image_id=f"img-{i}-0",
                alt_text=f"Alt text {i}.",
                confidence_score=0.9,
                validation_passed=True,
                validation_warnings=[],
                tokens_used=100,
                processing_time_seconds=1.0,
                timestamp=datetime.now(),
            )
            for i in range(3)
        ]

        # Mock internal method
        assembler._apply_alt_text_to_image = Mock(
            return_value="success"
        )

        status_map = assembler.apply_alt_text(results)

        assert len(status_map) == 3
        assert all(s == "success" for s in status_map.values())

    def test_apply_alt_text_handles_errors(
        self,
        sample_docx: Path,
        output_path: Path,
        sample_alt_text_result: AltTextResult
    ):
        """Test error handling during alt-text application."""
        assembler = DOCXAssembler(sample_docx, output_path)

        # Mock internal method to raise exception
        assembler._apply_alt_text_to_image = Mock(
            side_effect=Exception("Test error")
        )

        status_map = assembler.apply_alt_text([sample_alt_text_result])

        assert len(status_map) == 1
        assert "failed" in status_map[sample_alt_text_result.image_id]
        assert "test error" in status_map[
            sample_alt_text_result.image_id
        ].lower()

    def test_apply_alt_text_empty_list(
        self,
        sample_docx: Path,
        output_path: Path
    ):
        """Test alt-text application with empty result list."""
        assembler = DOCXAssembler(sample_docx, output_path)

        status_map = assembler.apply_alt_text([])

        assert len(status_map) == 0


class TestDOCXAssemblerImageMatching:
    """Test image matching and identification."""

    def test_apply_alt_text_invalid_image_id_format(
        self,
        sample_docx: Path,
        output_path: Path
    ):
        """Test handling of invalid image_id format."""
        assembler = DOCXAssembler(sample_docx, output_path)

        result = AltTextResult(
            image_id="invalid-format",
            alt_text="Test alt text.",
            confidence_score=0.9,
            validation_passed=True,
            validation_warnings=[],
            tokens_used=100,
            processing_time_seconds=1.0,
            timestamp=datetime.now(),
        )

        status = assembler._apply_alt_text_to_image(result)

        assert "invalid image_id format" in status.lower()

    def test_apply_alt_text_paragraph_out_of_range(
        self,
        sample_docx: Path,
        output_path: Path
    ):
        """Test handling of paragraph index out of range."""
        assembler = DOCXAssembler(sample_docx, output_path)

        result = AltTextResult(
            image_id="img-999-0",  # Non-existent paragraph
            alt_text="Test alt text.",
            confidence_score=0.9,
            validation_passed=True,
            validation_warnings=[],
            tokens_used=100,
            processing_time_seconds=1.0,
            timestamp=datetime.now(),
        )

        status = assembler._apply_alt_text_to_image(result)

        assert "out of range" in status.lower()

    def test_find_images_no_images_in_paragraph(
        self,
        sample_docx: Path,
        output_path: Path
    ):
        """Test finding images in paragraph with no images."""
        assembler = DOCXAssembler(sample_docx, output_path)

        paragraph = assembler.document.paragraphs[0]
        images = assembler._find_images_in_paragraph(paragraph)

        assert len(images) == 0


class TestDOCXAssemblerSaveDocument:
    """Test document saving functionality."""

    def test_save_document_success(
        self,
        sample_docx: Path,
        output_path: Path
    ):
        """Test successful document save."""
        assembler = DOCXAssembler(sample_docx, output_path)

        # Add a paragraph to modify the document
        assembler.document.add_paragraph("Modified content")

        assembler.save_document()

        assert output_path.exists()
        assert output_path.stat().st_size > 0

        # Verify it's a valid DOCX
        loaded_doc = Document(str(output_path))
        assert len(loaded_doc.paragraphs) > 0

    def test_save_document_preserves_content(
        self,
        sample_docx: Path,
        output_path: Path
    ):
        """Test saved document preserves original content."""
        assembler = DOCXAssembler(sample_docx, output_path)

        original_text = assembler.document.paragraphs[0].text

        assembler.save_document()

        # Load and verify
        loaded_doc = Document(str(output_path))
        assert loaded_doc.paragraphs[0].text == original_text

    def test_save_document_failure(
        self,
        sample_docx: Path,
        tmp_path: Path
    ):
        """Test save failure handling."""
        # Create output path in non-writable location (simulated)
        output_path = tmp_path / "output.docx"

        assembler = DOCXAssembler(sample_docx, output_path)

        # Mock save to raise exception
        assembler.document.save = Mock(
            side_effect=PermissionError("No permission")
        )

        with pytest.raises(ProcessingError) as exc_info:
            assembler.save_document()

        assert "failed to save" in str(exc_info.value).lower()


class TestDOCXAssemblerValidation:
    """Test document validation."""

    def test_validate_document_success(
        self,
        sample_docx: Path,
        output_path: Path
    ):
        """Test validation of valid document."""
        assembler = DOCXAssembler(sample_docx, output_path)

        assert assembler.validate_document() is True

    def test_validate_document_empty_file(
        self,
        tmp_path: Path,
        output_path: Path
    ):
        """Test validation of empty document."""
        empty_docx = tmp_path / "empty.docx"
        empty_docx.write_bytes(b"")

        # This should fail during initialization
        with pytest.raises(ProcessingError):
            DOCXAssembler(empty_docx, output_path)


class TestDOCXAssemblerIntegration:
    """Integration tests for complete workflow."""

    def test_complete_workflow(
        self,
        sample_docx: Path,
        output_path: Path,
        sample_alt_text_result: AltTextResult
    ):
        """Test complete workflow from load to save."""
        assembler = DOCXAssembler(sample_docx, output_path)

        # Validate
        assert assembler.validate_document() is True

        # Apply alt-text (will fail gracefully since no images)
        status_map = assembler.apply_alt_text([sample_alt_text_result])

        # Save
        assembler.save_document()

        # Verify output exists and is valid
        assert output_path.exists()

        loaded_doc = Document(str(output_path))
        assert len(loaded_doc.paragraphs) > 0

    def test_workflow_with_multiple_operations(
        self,
        sample_docx: Path,
        output_path: Path
    ):
        """Test workflow with multiple operations."""
        assembler = DOCXAssembler(sample_docx, output_path)

        # Add content
        assembler.document.add_paragraph("Additional paragraph")

        # Apply alt-text (empty list)
        status_map = assembler.apply_alt_text([])
        assert len(status_map) == 0

        # Save
        assembler.save_document()

        # Verify
        assert output_path.exists()

        loaded_doc = Document(str(output_path))
        assert len(loaded_doc.paragraphs) == 2  # Original + added
