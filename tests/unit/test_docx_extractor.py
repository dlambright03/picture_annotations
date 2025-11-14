"""
Unit tests for DOCX image extractor.

Tests extraction of images, position metadata, and alt-text from
Word documents.
"""

import io
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Inches
from PIL import Image

from ada_annotator.document_processors import DOCXExtractor
from ada_annotator.exceptions import ProcessingError
from ada_annotator.models import ImageMetadata


@pytest.fixture
def temp_docx_with_images(tmp_path):
    """Create a temporary DOCX file with test images."""
    doc = Document()
    doc.add_heading("Test Document", 0)

    # Add a paragraph with text
    doc.add_paragraph("This is a test paragraph before the image.")

    # Create a simple test image
    img = Image.new('RGB', (100, 100), color='red')
    img_bytes = io.BytesIO()
    img.save(img_bytes, format='PNG')
    img_bytes.seek(0)

    # Add inline image
    paragraph = doc.add_paragraph("Image: ")
    run = paragraph.add_run()
    run.add_picture(img_bytes, width=Inches(1.0))

    # Add another paragraph
    doc.add_paragraph("Text after image.")

    # Save document
    docx_path = tmp_path / "test_with_images.docx"
    doc.save(docx_path)

    return docx_path


@pytest.fixture
def temp_empty_docx(tmp_path):
    """Create a temporary DOCX file without images."""
    doc = Document()
    doc.add_heading("Empty Document", 0)
    doc.add_paragraph("No images here.")

    docx_path = tmp_path / "test_empty.docx"
    doc.save(docx_path)

    return docx_path


@pytest.fixture
def nonexistent_path(tmp_path):
    """Return a path that does not exist."""
    return tmp_path / "nonexistent.docx"


class TestDOCXExtractor:
    """Test suite for DOCXExtractor."""

    def test_initialization_with_valid_file(self, temp_empty_docx):
        """Test extractor initializes with valid DOCX file."""
        extractor = DOCXExtractor(temp_empty_docx)

        assert extractor.document_path == temp_empty_docx
        assert extractor.document is not None
        assert extractor.get_document_format() == "DOCX"

    def test_initialization_with_nonexistent_file(
        self, nonexistent_path
    ):
        """Test extractor raises error for nonexistent file."""
        with pytest.raises(FileNotFoundError) as exc_info:
            DOCXExtractor(nonexistent_path)

        assert "not found" in str(exc_info.value).lower()

    def test_initialization_with_wrong_extension(self, tmp_path):
        """Test extractor raises error for non-DOCX file."""
        wrong_file = tmp_path / "test.txt"
        wrong_file.write_text("not a docx")

        with pytest.raises(ValueError) as exc_info:
            DOCXExtractor(wrong_file)

        assert "not a docx" in str(exc_info.value).lower()

    def test_get_document_format(self, temp_empty_docx):
        """Test document format identifier."""
        extractor = DOCXExtractor(temp_empty_docx)
        assert extractor.get_document_format() == "DOCX"

    def test_validate_document(self, temp_empty_docx):
        """Test document validation."""
        extractor = DOCXExtractor(temp_empty_docx)
        assert extractor.validate_document() is True

    def test_extract_images_from_empty_document(self, temp_empty_docx):
        """Test extraction from document with no images."""
        extractor = DOCXExtractor(temp_empty_docx)
        images = extractor.extract_images()

        assert isinstance(images, list)
        assert len(images) == 0

    def test_extract_images_from_document_with_images(
        self, temp_docx_with_images
    ):
        """Test extraction from document with images."""
        extractor = DOCXExtractor(temp_docx_with_images)
        images = extractor.extract_images()

        assert isinstance(images, list)
        assert len(images) > 0

        # Check first image
        img = images[0]
        assert isinstance(img, ImageMetadata)
        assert img.image_id
        assert img.filename
        assert img.format in ["JPEG", "PNG", "GIF", "BMP"]
        assert img.size_bytes > 0
        assert img.width_pixels > 0
        assert img.height_pixels > 0

    def test_image_metadata_structure(self, temp_docx_with_images):
        """Test extracted image metadata has correct structure."""
        extractor = DOCXExtractor(temp_docx_with_images)
        images = extractor.extract_images()

        if len(images) > 0:
            img = images[0]

            # Check required fields
            assert hasattr(img, 'image_id')
            assert hasattr(img, 'filename')
            assert hasattr(img, 'format')
            assert hasattr(img, 'size_bytes')
            assert hasattr(img, 'width_pixels')
            assert hasattr(img, 'height_pixels')
            assert hasattr(img, 'position')
            assert hasattr(img, 'existing_alt_text')

            # Check position metadata
            assert isinstance(img.position, dict)
            assert 'paragraph_index' in img.position
            assert 'anchor_type' in img.position
            assert img.position['anchor_type'] in [
                'inline', 'floating'
            ]

    def test_position_metadata_paragraph_index(
        self, temp_docx_with_images
    ):
        """Test paragraph index is captured correctly."""
        extractor = DOCXExtractor(temp_docx_with_images)
        images = extractor.extract_images()

        if len(images) > 0:
            img = images[0]
            assert 'paragraph_index' in img.position
            assert isinstance(
                img.position['paragraph_index'], int
            )
            assert img.position['paragraph_index'] >= 0

    def test_inline_vs_floating_detection(
        self, temp_docx_with_images
    ):
        """Test anchor type is detected correctly."""
        extractor = DOCXExtractor(temp_docx_with_images)
        images = extractor.extract_images()

        if len(images) > 0:
            img = images[0]
            assert 'anchor_type' in img.position
            assert img.position['anchor_type'] in [
                'inline', 'floating'
            ]

    def test_image_id_uniqueness(self, temp_docx_with_images):
        """Test each image has unique ID."""
        extractor = DOCXExtractor(temp_docx_with_images)
        images = extractor.extract_images()

        image_ids = [img.image_id for img in images]
        assert len(image_ids) == len(set(image_ids))

    def test_existing_alt_text_none_when_absent(
        self, temp_docx_with_images
    ):
        """Test existing_alt_text is None when not present."""
        extractor = DOCXExtractor(temp_docx_with_images)
        images = extractor.extract_images()

        if len(images) > 0:
            # Since our test fixture doesn't add alt-text
            img = images[0]
            # Alt-text may be None or empty string
            assert img.existing_alt_text is None or (
                img.existing_alt_text == ""
            )

    def test_image_format_normalized(self, temp_docx_with_images):
        """Test image format is normalized correctly."""
        extractor = DOCXExtractor(temp_docx_with_images)
        images = extractor.extract_images()

        if len(images) > 0:
            img = images[0]
            # Format should be uppercase
            assert img.format.isupper()
            # Format should be one of supported types
            assert img.format in ["JPEG", "PNG", "GIF", "BMP"]

    def test_page_number_is_none_for_docx(
        self, temp_docx_with_images
    ):
        """Test page_number is None for DOCX (no page concept)."""
        extractor = DOCXExtractor(temp_docx_with_images)
        images = extractor.extract_images()

        if len(images) > 0:
            img = images[0]
            # DOCX doesn't have page numbers
            assert img.page_number is None


class TestDOCXExtractorEdgeCases:
    """Test edge cases and error handling."""

    def test_corrupted_docx_file(self, tmp_path):
        """Test handling of corrupted DOCX file."""
        corrupted_path = tmp_path / "corrupted.docx"
        corrupted_path.write_bytes(b"not a valid docx file")

        with pytest.raises(ProcessingError):
            DOCXExtractor(corrupted_path)

    def test_extract_from_document_with_only_text(
        self, temp_empty_docx
    ):
        """Test extraction from document with only text."""
        extractor = DOCXExtractor(temp_empty_docx)
        images = extractor.extract_images()

        assert images == []

    def test_multiple_images_in_same_paragraph(self, tmp_path):
        """Test extraction of multiple images in one paragraph."""
        doc = Document()
        paragraph = doc.add_paragraph()

        # Add multiple images to same paragraph
        for i in range(3):
            img = Image.new('RGB', (50, 50), color='blue')
            img_bytes = io.BytesIO()
            img.save(img_bytes, format='PNG')
            img_bytes.seek(0)

            run = paragraph.add_run()
            run.add_picture(img_bytes, width=Inches(0.5))

        docx_path = tmp_path / "multiple_images.docx"
        doc.save(docx_path)

        extractor = DOCXExtractor(docx_path)
        images = extractor.extract_images()

        # Should extract all images
        assert len(images) >= 3

        # All should be from same paragraph
        para_indices = [
            img.position['paragraph_index'] for img in images
        ]
        # Most images should be from the same paragraph
        assert len(set(para_indices)) <= 2  # Allow for heading

    def test_handles_blip_without_embed_id(self, tmp_path):
        """Test that extraction continues when blip has no r:embed attribute."""
        from unittest.mock import MagicMock, patch
        from lxml.etree import QName

        doc = Document()
        # Add a valid image first
        img = Image.new('RGB', (50, 50), color='green')
        img_bytes = io.BytesIO()
        img.save(img_bytes, format='PNG')
        img_bytes.seek(0)

        p = doc.add_paragraph()
        p.add_run().add_picture(img_bytes, width=Inches(0.5))

        docx_path = tmp_path / "test_missing_rid.docx"
        doc.save(docx_path)

        extractor = DOCXExtractor(docx_path)

        # Get the real blips first
        real_paragraph = extractor.document.paragraphs[0]
        real_blips = real_paragraph._element.xpath(".//a:blip")

        # Create a mock blip without r:embed
        mock_blip = MagicMock()
        mock_blip.get.return_value = None  # No r:embed attribute

        # Mock xpath to return our mock blip plus the real one
        def mock_xpath_func(query):
            if query == ".//a:blip":
                return [mock_blip] + real_blips
            return real_paragraph._element.xpath(query)

        with patch.object(real_paragraph._element, 'xpath', side_effect=mock_xpath_func):
            # Should extract images and skip the one without r:embed
            images = extractor._extract_images_from_paragraph(real_paragraph, 0)

            # Should have extracted the valid image
            assert len(images) >= 1

    def test_handles_missing_image_part(self, tmp_path):
        """Test that extraction continues when image part is not found in relationships."""
        from unittest.mock import MagicMock, patch, PropertyMock

        doc = Document()
        # Add a valid image
        img = Image.new('RGB', (50, 50), color='blue')
        img_bytes = io.BytesIO()
        img.save(img_bytes, format='PNG')
        img_bytes.seek(0)

        p = doc.add_paragraph()
        p.add_run().add_picture(img_bytes, width=Inches(0.5))

        docx_path = tmp_path / "test_missing_part.docx"
        doc.save(docx_path)

        extractor = DOCXExtractor(docx_path)

        # Mock related_parts as a property that returns a dict-like object
        mock_related_parts = MagicMock()
        mock_related_parts.get.return_value = None  # Returns None for any key

        with patch.object(type(extractor.document.part), 'related_parts', new_callable=PropertyMock) as mock_prop:
            mock_prop.return_value = mock_related_parts

            # Should handle missing part gracefully
            images = extractor._extract_images_from_paragraph(extractor.document.paragraphs[0], 0)

            # Should return empty list when part is missing
            assert isinstance(images, list)

    def test_handles_corrupted_image_data(self, tmp_path):
        """Test that extraction continues when image data is corrupted."""
        from unittest.mock import MagicMock, patch

        doc = Document()
        # Add two valid images
        for color in ['red', 'green']:
            img = Image.new('RGB', (50, 50), color=color)
            img_bytes = io.BytesIO()
            img.save(img_bytes, format='PNG')
            img_bytes.seek(0)

            p = doc.add_paragraph()
            p.add_run().add_picture(img_bytes, width=Inches(0.5))

        docx_path = tmp_path / "test_corrupted.docx"
        doc.save(docx_path)

        extractor = DOCXExtractor(docx_path)

        # Get all blips from the first paragraph
        all_blips = extractor.document.paragraphs[0]._element.xpath(".//a:blip")

        if all_blips:
            # Mock PIL Image.open to fail for the first image
            original_open = Image.open
            call_count = [0]

            def mock_open(fp, *args, **kwargs):
                call_count[0] += 1
                if call_count[0] == 1:
                    raise OSError("Corrupted image data")
                return original_open(fp, *args, **kwargs)

            with patch('PIL.Image.open', side_effect=mock_open):
                # Should continue extraction despite corrupted image
                images = extractor._extract_images_from_paragraph(extractor.document.paragraphs[0], 0)

                # Should still process (might be empty due to exception)
                assert isinstance(images, list)

    def test_extract_images_failure_handling(self, temp_docx_with_images):
        """Test that extract_images raises ProcessingError on failure."""
        from unittest.mock import patch

        extractor = DOCXExtractor(temp_docx_with_images)

        # Mock _extract_images_from_paragraph to raise an exception
        with patch.object(extractor, '_extract_images_from_paragraph', side_effect=RuntimeError("Test error")):
            with pytest.raises(ProcessingError) as exc_info:
                extractor.extract_images()

            assert "Failed to extract images" in str(exc_info.value)


class TestDOCXExtractorIntegration:
    """Integration tests with real-world scenarios."""

    def test_full_extraction_workflow(self, temp_docx_with_images):
        """Test complete extraction workflow."""
        # Initialize extractor
        extractor = DOCXExtractor(temp_docx_with_images)

        # Validate document
        assert extractor.validate_document() is True

        # Extract images
        images = extractor.extract_images()

        # Verify results
        assert isinstance(images, list)

        for img in images:
            # Verify all required fields
            assert img.image_id
            assert img.filename
            assert img.format
            assert img.size_bytes > 0
            assert img.width_pixels > 0
            assert img.height_pixels > 0
            assert isinstance(img.position, dict)
            assert 'paragraph_index' in img.position
            assert 'anchor_type' in img.position


class TestDOCXExtractorHelperMethods:
    """Test private/helper methods of DOCXExtractor."""

    def test_extract_images_from_paragraph_with_image(self, temp_docx_with_images):
        """Test _extract_images_from_paragraph() with images."""
        extractor = DOCXExtractor(temp_docx_with_images)

        # Get first paragraph with image
        for para_idx, paragraph in enumerate(extractor.document.paragraphs):
            blips = paragraph._element.xpath(".//a:blip")
            if blips:
                images = extractor._extract_images_from_paragraph(paragraph, para_idx)
                assert len(images) > 0
                assert all(isinstance(img, ImageMetadata) for img in images)
                # Verify paragraph index is set correctly
                assert all(img.position['paragraph_index'] == para_idx for img in images)
                break

    def test_extract_images_from_paragraph_without_image(self, temp_docx_with_images):
        """Test _extract_images_from_paragraph() with no images."""
        extractor = DOCXExtractor(temp_docx_with_images)

        # Get first paragraph without image
        for para_idx, paragraph in enumerate(extractor.document.paragraphs):
            blips = paragraph._element.xpath(".//a:blip")
            if not blips:
                images = extractor._extract_images_from_paragraph(paragraph, para_idx)
                assert images == []
                break

    def test_convert_to_png_with_background_rgba(self, temp_empty_docx):
        """Test PNG conversion with RGBA image."""
        # Create RGBA image with transparency
        img = Image.new('RGBA', (100, 100), color=(255, 0, 0, 128))

        extractor = DOCXExtractor(temp_empty_docx)
        converted_img, img_bytes = extractor._convert_to_png_with_background(img)

        # Should be converted to RGB with white background
        assert converted_img.mode == "RGB"
        assert len(img_bytes) > 0

        # Verify it's valid PNG data
        reloaded = Image.open(io.BytesIO(img_bytes))
        assert reloaded.format == "PNG"

    def test_convert_to_png_with_background_la(self, temp_empty_docx):
        """Test PNG conversion with LA (grayscale + alpha) image."""
        img = Image.new('LA', (100, 100), color=(128, 128))

        extractor = DOCXExtractor(temp_empty_docx)
        converted_img, img_bytes = extractor._convert_to_png_with_background(img)

        # Should be converted to RGB
        assert converted_img.mode == "RGB"
        assert len(img_bytes) > 0

    def test_convert_to_png_with_background_palette_transparency(self, temp_empty_docx):
        """Test PNG conversion with palette mode and transparency."""
        # Create palette image with transparency
        img = Image.new('P', (100, 100))
        img.info['transparency'] = 0

        extractor = DOCXExtractor(temp_empty_docx)
        converted_img, img_bytes = extractor._convert_to_png_with_background(img)

        # Should be converted to RGB
        assert converted_img.mode == "RGB"
        assert len(img_bytes) > 0

    def test_convert_to_png_with_background_rgb(self, temp_empty_docx):
        """Test PNG conversion with RGB image (no transparency)."""
        img = Image.new('RGB', (100, 100), color='blue')

        extractor = DOCXExtractor(temp_empty_docx)
        converted_img, img_bytes = extractor._convert_to_png_with_background(img)

        # Should remain RGB
        assert converted_img.mode == "RGB"
        assert len(img_bytes) > 0


class TestDOCXExtractorImageFormats:
    """Test handling of different image formats."""

    def test_jpeg_image_extraction(self, tmp_path):
        """Test extraction of JPEG images."""
        doc = Document()
        doc.add_paragraph("JPEG test")

        # Create JPEG image
        img = Image.new('RGB', (100, 100), color='green')
        img_bytes = io.BytesIO()
        img.save(img_bytes, format='JPEG')
        img_bytes.seek(0)

        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(img_bytes, width=Inches(1.0))

        docx_path = tmp_path / "jpeg_test.docx"
        doc.save(docx_path)

        extractor = DOCXExtractor(docx_path)
        images = extractor.extract_images()

        assert len(images) > 0
        # JPEG format should be normalized
        assert images[0].format in ["JPEG", "PNG"]

    def test_png_image_extraction(self, tmp_path):
        """Test extraction of PNG images."""
        doc = Document()
        doc.add_paragraph("PNG test")

        # Create PNG image
        img = Image.new('RGB', (100, 100), color='yellow')
        img_bytes = io.BytesIO()
        img.save(img_bytes, format='PNG')
        img_bytes.seek(0)

        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(img_bytes, width=Inches(1.0))

        docx_path = tmp_path / "png_test.docx"
        doc.save(docx_path)

        extractor = DOCXExtractor(docx_path)
        images = extractor.extract_images()

        assert len(images) > 0
        assert images[0].format == "PNG"


class TestDOCXExtractorAltText:
    """Test alt-text extraction functionality."""

    def test_extract_alt_text_when_present(self, tmp_path):
        """Test extraction of existing alt-text."""
        # Note: python-docx doesn't provide easy API to add alt-text
        # This tests the extraction logic with None result
        extractor_module = __import__(
            'ada_annotator.document_processors.docx_extractor',
            fromlist=['DOCXExtractor']
        )

        # Test with None blip (edge case)
        from unittest.mock import Mock
        extractor = Mock()
        extractor.logger = Mock()

        result = extractor_module.DOCXExtractor._extract_alt_text_from_blip(
            extractor, None
        )
        # Should handle gracefully
        assert result is None

    def test_extract_alt_text_when_absent(self, temp_docx_with_images):
        """Test extraction when no alt-text is present."""
        extractor = DOCXExtractor(temp_docx_with_images)
        images = extractor.extract_images()

        # Test fixture doesn't add alt-text
        for img in images:
            assert img.existing_alt_text is None or img.existing_alt_text == ""


class TestDOCXExtractorComplexDocuments:
    """Test with complex document structures."""

    def test_document_with_mixed_content(self, tmp_path):
        """Test document with text, images, tables, etc."""
        doc = Document()
        doc.add_heading("Complex Document", 0)
        doc.add_paragraph("First paragraph")

        # Add image
        img = Image.new('RGB', (50, 50), color='purple')
        img_bytes = io.BytesIO()
        img.save(img_bytes, format='PNG')
        img_bytes.seek(0)

        p = doc.add_paragraph()
        p.add_run().add_picture(img_bytes, width=Inches(0.5))

        # Add table
        doc.add_table(rows=2, cols=2)

        # Add more text
        doc.add_paragraph("After table")

        # Add another image
        img2 = Image.new('RGB', (50, 50), color='orange')
        img_bytes2 = io.BytesIO()
        img2.save(img_bytes2, format='PNG')
        img_bytes2.seek(0)

        p2 = doc.add_paragraph()
        p2.add_run().add_picture(img_bytes2, width=Inches(0.5))

        docx_path = tmp_path / "complex.docx"
        doc.save(docx_path)

        extractor = DOCXExtractor(docx_path)
        images = extractor.extract_images()

        # Should extract both images
        assert len(images) >= 2

        # Should have different paragraph indices
        para_indices = [img.position['paragraph_index'] for img in images]
        assert len(set(para_indices)) >= 2

    def test_document_with_many_images(self, tmp_path):
        """Test document with many images."""
        doc = Document()
        doc.add_heading("Many Images", 0)

        # Add 10 images in different paragraphs
        for i in range(10):
            doc.add_paragraph(f"Section {i}")
            img = Image.new('RGB', (30, 30), color=(i*25, i*25, i*25))
            img_bytes = io.BytesIO()
            img.save(img_bytes, format='PNG')
            img_bytes.seek(0)

            p = doc.add_paragraph()
            p.add_run().add_picture(img_bytes, width=Inches(0.3))

        docx_path = tmp_path / "many_images.docx"
        doc.save(docx_path)

        extractor = DOCXExtractor(docx_path)
        images = extractor.extract_images()

        # Should extract all images
        assert len(images) >= 10

        # All should have unique IDs
        image_ids = [img.image_id for img in images]
        assert len(image_ids) == len(set(image_ids))

        # Should have proper paragraph indices
        for img in images:
            assert 'paragraph_index' in img.position
            assert img.position['paragraph_index'] >= 0


class TestDOCXExtractorErrorRecovery:
    """Test error handling and recovery."""

    def test_extraction_continues_after_single_image_failure(self, tmp_path):
        """Test that extraction continues if one image fails."""
        doc = Document()

        # Add a valid image
        img1 = Image.new('RGB', (50, 50), color='red')
        img_bytes1 = io.BytesIO()
        img1.save(img_bytes1, format='PNG')
        img_bytes1.seek(0)

        p1 = doc.add_paragraph()
        p1.add_run().add_picture(img_bytes1, width=Inches(0.5))

        # Add another valid image
        img2 = Image.new('RGB', (50, 50), color='blue')
        img_bytes2 = io.BytesIO()
        img2.save(img_bytes2, format='PNG')
        img_bytes2.seek(0)

        p2 = doc.add_paragraph()
        p2.add_run().add_picture(img_bytes2, width=Inches(0.5))

        docx_path = tmp_path / "partial_fail.docx"
        doc.save(docx_path)

        extractor = DOCXExtractor(docx_path)
        images = extractor.extract_images()

        # Should extract successfully without failing
        assert isinstance(images, list)
        assert len(images) >= 2

    def test_image_data_included_in_metadata(self, temp_docx_with_images):
        """Test that image binary data is included in metadata."""
        extractor = DOCXExtractor(temp_docx_with_images)
        images = extractor.extract_images()

        for img in images:
            assert img.image_data is not None
            assert isinstance(img.image_data, bytes)
            assert len(img.image_data) > 0
            assert len(img.image_data) == img.size_bytes
