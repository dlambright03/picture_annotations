"""
DOCX document image extractor.

Extracts images from Word documents with position metadata and
existing alt-text. Handles both inline and floating images.
"""

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from PIL import Image

from ada_annotator.document_processors.base_extractor import DocumentExtractor
from ada_annotator.exceptions import ProcessingError
from ada_annotator.models import ImageMetadata


class DOCXExtractor(DocumentExtractor):
    """
    Extract images from DOCX documents.

    Supports:
    - Inline images (in text flow)
    - Floating/anchored images
    - Position metadata (paragraph index)
    - Existing alt-text extraction
    """

    def __init__(self, document_path: Path):
        """
        Initialize DOCX extractor.

        Args:
            document_path: Path to DOCX file.

        Raises:
            FileNotFoundError: If file does not exist.
            ValueError: If file is not DOCX format.
        """
        super().__init__(document_path)

        # Validate DOCX extension
        if document_path.suffix.lower() != ".docx":
            raise ValueError(f"Not a DOCX file: {document_path}")

        # Load document
        try:
            self.document = Document(document_path)
            self.logger.info(
                "docx_loaded",
                document_path=str(document_path),
                paragraph_count=len(self.document.paragraphs),
            )
        except Exception as e:
            raise ProcessingError(f"Failed to load DOCX: {e}") from e

    def get_document_format(self) -> str:
        """
        Get document format identifier.

        Returns:
            str: 'DOCX'
        """
        return "DOCX"

    def extract_images(self) -> list[ImageMetadata]:
        """
        Extract all images from DOCX document.

        Extracts both inline and floating images with position
        metadata and existing alt-text.

        Returns:
            List[ImageMetadata]: List of extracted images.

        Raises:
            ProcessingError: If extraction fails.
        """
        self.logger.info(
            "extraction_started",
            document_format="DOCX",
        )

        images = []

        try:
            # Extract all images paragraph by paragraph to maintain consistent indexing
            for para_idx, paragraph in enumerate(self.document.paragraphs):
                para_images = self._extract_images_from_paragraph(paragraph, para_idx)
                images.extend(para_images)

            self.logger.info(
                "extraction_completed",
                total_images=len(images),
            )

            return images

        except Exception as e:
            self.logger.error(
                "extraction_failed",
                error=str(e),
            )
            raise ProcessingError(f"Failed to extract images: {e}") from e

    def _extract_images_from_paragraph(
        self, paragraph, para_idx: int
    ) -> list[ImageMetadata]:
        """
        Extract all images (inline and floating) from a single paragraph.

        Args:
            paragraph: The paragraph object.
            para_idx: The paragraph index.

        Returns:
            List[ImageMetadata]: Images found in this paragraph.
        """
        images = []

        # Find all blips (images) in this paragraph (both inline and floating)
        all_blips = paragraph._element.xpath(".//a:blip")

        for blip in all_blips:
            try:
                # Extract image from relationship
                rId = blip.get(qn("r:embed"))
                if not rId:
                    continue

                image_part = self.document.part.related_parts.get(rId)
                if not image_part:
                    continue

                # Get image binary data
                image_bytes = image_part.blob

                # Get image format from content type
                content_type = image_part.content_type
                format_str = content_type.split("/")[-1].upper()

                # Normalize format names
                if format_str in ["JPG", "JPEG"]:
                    format_str = "JPEG"

                # Load with PIL to get dimensions
                img = Image.open(BytesIO(image_bytes))

                # Convert unsupported formats (EMF, WMF, etc.) to PNG
                # Also add white background to transparent images
                if format_str not in ["JPEG", "PNG", "GIF", "BMP"]:
                    self.logger.debug(
                        "converting_unsupported_format",
                        paragraph_index=para_idx,
                        original_format=format_str,
                        target_format="PNG",
                    )
                    img, image_bytes = self._convert_to_png_with_background(img)
                    format_str = "PNG"
                elif img.mode in ("RGBA", "LA", "PA") or (
                    img.mode == "P" and "transparency" in img.info
                ):
                    # Image has transparency - add white background
                    self.logger.debug(
                        "adding_white_background",
                        paragraph_index=para_idx,
                        image_mode=img.mode,
                    )
                    img, image_bytes = self._convert_to_png_with_background(img)
                    format_str = "PNG"

                # Extract existing alt-text
                alt_text = self._extract_alt_text_from_blip(blip)

                # Determine anchor type (inline vs floating)
                anchor_type = "inline"
                if blip.xpath("ancestor::w:drawing/wp:anchor"):
                    anchor_type = "floating"

                # Create position metadata
                position = {
                    "paragraph_index": para_idx,
                    "anchor_type": anchor_type,
                }

                # Create ImageMetadata with per-paragraph image index
                image_idx = len(images)  # Index within this paragraph
                image_id = f"img-{para_idx}-{image_idx}"
                metadata = ImageMetadata(
                    image_id=image_id,
                    filename=(f"{image_id}.{format_str.lower()}"),
                    format=format_str,
                    size_bytes=len(image_bytes),
                    width_pixels=img.width,
                    height_pixels=img.height,
                    page_number=None,  # DOCX has no pages
                    position=position,
                    existing_alt_text=alt_text,
                    image_data=image_bytes,  # Include binary data
                )

                images.append(metadata)

                self.logger.debug(
                    "image_extracted",
                    image_id=image_id,
                    paragraph_index=para_idx,
                    format=format_str,
                    anchor_type=anchor_type,
                    has_alt_text=bool(alt_text),
                )

            except Exception as e:
                self.logger.warning(
                    "image_extraction_failed",
                    paragraph_index=para_idx,
                    error=str(e),
                )
                continue

        return images

    def _extract_alt_text_from_blip(self, blip) -> str | None:
        """
        Extract existing alt-text from image element.

        Args:
            blip: The blip (image) element.

        Returns:
            Optional[str]: Alt-text if present, None otherwise.
        """
        try:
            # Try to find docPr (document properties) element
            # which contains title/description
            doc_pr_elements = blip.xpath(
                "ancestor::wp:inline/wp:docPr | "
                "ancestor::wp:anchor/wp:docPr",
                namespaces={
                    "wp": (
                        "http://schemas.openxmlformats.org/"
                        "drawingml/2006/wordprocessingDrawing"
                    )
                },
            )

            for doc_pr in doc_pr_elements:
                # Check for title attribute
                title = doc_pr.get("title")
                if title:
                    return title

                # Check for descr attribute
                descr = doc_pr.get("descr")
                if descr:
                    return descr

            return None

        except Exception:
            return None

    def _convert_to_png_with_background(self, img: Image.Image) -> tuple[Image.Image, bytes]:
        """
        Convert image to PNG format with white background.

        Handles transparent images by compositing them onto a white background.
        Also converts unsupported formats (EMF, WMF, etc.) to PNG.

        Args:
            img: PIL Image object.

        Returns:
            tuple: (converted Image object, PNG bytes)
        """
        # Create a white background image
        if img.mode in ("RGBA", "LA"):
            # Has alpha channel
            background = Image.new("RGB", img.size, (255, 255, 255))
            if img.mode == "RGBA":
                background.paste(img, mask=img.split()[3])  # Use alpha channel as mask
            else:  # LA mode
                background.paste(img.convert("RGB"), mask=img.split()[1])
            img = background
        elif img.mode == "PA":
            # Palette with alpha
            img = img.convert("RGBA")
            background = Image.new("RGB", img.size, (255, 255, 255))
            background.paste(img, mask=img.split()[3])
            img = background
        elif img.mode == "P" and "transparency" in img.info:
            # Palette with transparency
            img = img.convert("RGBA")
            background = Image.new("RGB", img.size, (255, 255, 255))
            background.paste(img, mask=img.split()[3])
            img = background
        elif img.mode not in ("RGB", "L"):
            # Convert other modes to RGB
            img = img.convert("RGB")

        # Save as PNG
        buffer = BytesIO()
        img.save(buffer, format="PNG")
        image_bytes = buffer.getvalue()

        return img, image_bytes
