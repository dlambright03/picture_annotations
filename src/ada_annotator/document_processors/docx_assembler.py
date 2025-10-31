"""
DOCX document assembler.

Applies AI-generated alt-text to images in Word documents while
preserving exact image positions and document structure.
"""

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from ada_annotator.document_processors.base_assembler import DocumentAssembler
from ada_annotator.exceptions import ProcessingError
from ada_annotator.models import AltTextResult


class DOCXAssembler(DocumentAssembler):
    """
    Apply alt-text to DOCX documents.

    Supports:
    - Inline images (in text flow)
    - Floating/anchored images
    - Position preservation (paragraph-based)
    - Existing alt-text handling
    """

    def __init__(self, input_path: Path, output_path: Path):
        """
        Initialize DOCX assembler.

        Args:
            input_path: Path to input DOCX file.
            output_path: Path where output DOCX will be saved.

        Raises:
            FileNotFoundError: If input file does not exist.
            ValueError: If file is not DOCX format.
        """
        super().__init__(input_path, output_path)

        # Validate DOCX extension
        if input_path.suffix.lower() != ".docx":
            raise ValueError(f"Not a DOCX file: {input_path}")

        try:
            self.document = Document(str(input_path))
            self.logger.info(
                "docx_loaded",
                paragraph_count=len(self.document.paragraphs),
                section_count=len(self.document.sections),
            )
        except Exception as e:
            raise ProcessingError(f"Failed to load DOCX document: {e}") from e

    def apply_alt_text(
        self, alt_text_results: list[AltTextResult]
    ) -> dict[str, str]:
        """
        Apply alt-text to images in DOCX document.

        Finds images by position metadata (paragraph index) and
        applies alt-text via XML manipulation. Preserves image
        positions and handles errors gracefully.

        Args:
            alt_text_results: List of alt-text generation results.

        Returns:
            Dict[str, str]: Map of image_id to status message.

        Raises:
            ProcessingError: If critical errors occur.
        """
        status_map = {}

        self.logger.info(
            "applying_alt_text",
            total_results=len(alt_text_results),
        )

        for result in alt_text_results:
            try:
                status = self._apply_alt_text_to_image(result)
                status_map[result.image_id] = status

                self.logger.debug(
                    "alt_text_applied",
                    image_id=result.image_id,
                    status=status,
                    alt_text_length=len(result.alt_text),
                )
            except Exception as e:
                status_map[result.image_id] = f"failed: {str(e)}"
                self.logger.warning(
                    "alt_text_application_failed",
                    image_id=result.image_id,
                    error=str(e),
                )

        success_count = sum(1 for s in status_map.values() if s == "success")
        self.logger.info(
            "alt_text_application_complete",
            success_count=success_count,
            total_count=len(alt_text_results),
        )

        return status_map

    def _apply_alt_text_to_image(self, result: AltTextResult) -> str:
        """
        Apply alt-text to a single image.

        Args:
            result: Alt-text generation result.

        Returns:
            str: Status message ('success', 'skipped', or error).
        """
        # Extract paragraph index and image index from image_id
        # Format: "img-{paragraph_idx}-{image_idx}"
        try:
            parts = result.image_id.split("-")
            paragraph_idx = int(parts[1])
            image_idx = int(parts[2])
        except (IndexError, ValueError):
            return "failed: invalid image_id format"

        # Validate paragraph index
        if paragraph_idx >= len(self.document.paragraphs):
            return "failed: paragraph index out of range"

        paragraph = self.document.paragraphs[paragraph_idx]

        # Find images in this paragraph
        images_found = self._find_images_in_paragraph(paragraph)

        if not images_found:
            return "failed: no images found in paragraph"

        # Validate image index
        if image_idx >= len(images_found):
            return "failed: image index out of range"

        # Apply alt-text to the specific image at the given index
        img_element = images_found[image_idx]
        
        # For decorative images, set empty alt-text
        alt_text_to_apply = "" if result.is_decorative else result.alt_text
        
        if self._set_alt_text_on_element(img_element, alt_text_to_apply):
            if result.is_decorative:
                return "success (decorative)"
            return "success"
        else:
            return "failed: could not set alt-text"

    def _find_images_in_paragraph(self, paragraph) -> list:
        """
        Find all image elements in a paragraph.

        Args:
            paragraph: docx.text.paragraph.Paragraph object.

        Returns:
            List: List of image XML elements (deduplicated).
        """
        images = []
        seen_elements = set()

        # Find inline images (in runs)
        for run in paragraph.runs:
            # Look for inline shapes (pictures)
            # Use qn() for qualified names
            inline_shapes = run._element.findall(f'.//{qn("pic:pic")}')
            for shape in inline_shapes:
                # Use id() to deduplicate - same object won't be added twice
                element_id = id(shape)
                if element_id not in seen_elements:
                    seen_elements.add(element_id)
                    images.append(shape)

        # Find floating images (anchored to paragraph)
        # Note: This may find images not in runs, so we deduplicate
        floating_shapes = paragraph._element.findall(f'.//{qn("pic:pic")}')
        for shape in floating_shapes:
            element_id = id(shape)
            if element_id not in seen_elements:
                seen_elements.add(element_id)
                images.append(shape)

        return images

    def _set_alt_text_on_element(self, img_element, alt_text: str) -> bool:
        """
        Set alt-text on an image XML element.

        Args:
            img_element: XML element for the image.
            alt_text: Alt-text to apply.

        Returns:
            bool: True if successful, False otherwise.
        """
        try:
            # Find cNvPr (non-visual properties) element
            # This contains title and descr attributes
            nvpicpr = img_element.find(
                qn("pic:nvPicPr"), namespaces=img_element.nsmap
            )

            if nvpicpr is None:
                self.logger.warning(
                    "no_nvpicpr_element",
                    message="Could not find nvPicPr element",
                )
                return False

            cnvpr = nvpicpr.find(qn("pic:cNvPr"), namespaces=img_element.nsmap)

            if cnvpr is None:
                self.logger.warning(
                    "no_cnvpr_element", message="Could not find cNvPr element"
                )
                return False

            # Set both title and descr attributes
            # (Word uses descr, some tools use title)
            cnvpr.set("title", alt_text)
            cnvpr.set("descr", alt_text)

            return True

        except Exception as e:
            self.logger.warning("set_alt_text_failed", error=str(e))
            return False

    def save_document(self) -> None:
        """
        Save modified document to output path.

        Raises:
            ProcessingError: If save operation fails.
        """
        try:
            self.document.save(str(self.output_path))

            self.logger.info(
                "document_saved",
                output_path=str(self.output_path),
                file_size_bytes=self.output_path.stat().st_size,
            )
        except Exception as e:
            raise ProcessingError(f"Failed to save DOCX document: {e}") from e

    def get_document_format(self) -> str:
        """
        Get document format identifier.

        Returns:
            str: 'DOCX'
        """
        return "DOCX"
