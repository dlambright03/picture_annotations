"""
Debug document generator for visual verification of alt-text.

Creates formatted DOCX documents showing extracted images alongside
their generated alt-text annotations.
"""

from io import BytesIO
from pathlib import Path

import structlog
from docx import Document
from docx.shared import Inches

from ada_annotator.models import AltTextResult, ImageMetadata


def create_debug_document(
    images: list[ImageMetadata],
    alt_text_results: list[AltTextResult],
    output_path: Path,
) -> None:
    """
    Create debug document with images and their alt-text.

    Generates a formatted DOCX document displaying each extracted image
    alongside its metadata and generated alt-text for visual verification.

    Args:
        images: List of extracted image metadata with binary data.
        alt_text_results: List of generated alt-text results.
        output_path: Path where debug document will be saved.

    Raises:
        OSError: If unable to save document to output path.
    """
    logger = structlog.get_logger(__name__)

    doc = Document()
    doc.add_heading("Alt-Text Debug Output", level=1)
    doc.add_paragraph(
        f"Total Images: {len(images)} | "
        f"Successful Annotations: {len(alt_text_results)}"
    )
    doc.add_paragraph()  # Spacing

    # Create lookup for alt-text by image_id
    alt_text_map = {r.image_id: r for r in alt_text_results}

    for idx, img in enumerate(images, 1):
        # Add image heading
        doc.add_heading(f"Image {idx}: {img.image_id}", level=2)

        # Add the actual image
        if img.image_data:
            try:
                image_stream = BytesIO(img.image_data)
                doc.add_picture(image_stream, width=Inches(4.0))
            except Exception as e:
                logger.warning(
                    "failed_to_add_image",
                    image_id=img.image_id,
                    error=str(e),
                )
                doc.add_paragraph(f"[Image could not be displayed: {e}]")
        else:
            doc.add_paragraph("[No image data available]")

        # Add metadata
        doc.add_paragraph(f"Filename: {img.filename}")
        doc.add_paragraph(f"Format: {img.format}")
        doc.add_paragraph(
            f"Dimensions: {img.width_pixels}x{img.height_pixels} pixels"
        )

        # Add generated alt-text
        if img.image_id in alt_text_map:
            result = alt_text_map[img.image_id]
            if result.is_decorative:
                doc.add_paragraph("Alt-Text: [DECORATIVE - No alt-text needed]")
                para = doc.add_paragraph()
                para.add_run("Status: ").bold = True
                para.add_run("Marked as decorative element")
            else:
                doc.add_paragraph(f"Alt-Text: {result.alt_text}")
                doc.add_paragraph(f"Confidence: {result.confidence_score:.2%}")
            doc.add_paragraph(f"Tokens Used: {result.tokens_used}")
            doc.add_paragraph(
                f"Processing Time: {result.processing_time_seconds:.2f}s"
            )
        else:
            doc.add_paragraph(
                "Alt-Text: [Generation failed or not processed]"
            )

        # Add spacing between images
        doc.add_paragraph()
        doc.add_paragraph("─" * 70)  # Visual separator
        doc.add_paragraph()

    # Save document
    doc.save(str(output_path))
    logger.info(
        "debug_document_created",
        output_path=str(output_path),
        total_images=len(images),
    )
